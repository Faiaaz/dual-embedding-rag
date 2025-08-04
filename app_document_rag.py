from flask import Flask, request, jsonify
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
import json
import os
import io
import fitz  # PyMuPDF for PDFs
from docx import Document  # python-docx for DOCX files
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
from typing import List, Dict, Any
import tempfile
from werkzeug.utils import secure_filename

app = Flask(__name__)

# Configure upload settings
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'csv', 'db'}

# Create upload folder if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max file size

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(pdf_path: str, use_ocr: bool = False) -> str:
    """Extract text from PDF using PyMuPDF with optional OCR support"""
    try:
        # Open the PDF
        doc = fitz.open(pdf_path)
        text = ""
        
        # Extract text from each page
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            
            # First try to extract text normally
            page_text = page.get_text()
            
            # If no text found and OCR is enabled and available, try OCR
            if (not page_text.strip() and use_ocr and OCR_AVAILABLE):
                try:
                    # Convert page to image
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # Higher resolution for better OCR
                    img_data = pix.tobytes("png")
                    
                    # Convert to PIL Image
                    img = Image.open(io.BytesIO(img_data))
                    
                    # Perform OCR
                    ocr_text = pytesseract.image_to_string(img)
                    page_text = ocr_text
                    
                except Exception as ocr_error:
                    print(f"OCR failed for page {page_num + 1}: {ocr_error}")
                    page_text = f"[OCR failed for page {page_num + 1}]"
            
            text += page_text + "\n"
        
        doc.close()
        return text.strip()
    
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from DOCX using python-docx"""
    try:
        # Open the DOCX file
        doc = Document(docx_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")

def extract_text_from_txt(txt_path: str) -> str:
    """Extract text from TXT file"""
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(txt_path, 'r', encoding=encoding) as file:
                    text = file.read()
                    return text.strip()
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, try with error handling
        with open(txt_path, 'r', encoding='utf-8', errors='ignore') as file:
            text = file.read()
            return text.strip()
    
    except Exception as e:
        raise Exception(f"Error extracting text from TXT: {str(e)}")

def extract_text_from_csv(csv_path: str) -> str:
    """Extract text from CSV file using pandas with enhanced formatting"""
    if not PANDAS_AVAILABLE:
        raise Exception("Pandas is not available. Please install pandas to process CSV files.")
    
    try:
        # Read CSV file with UTF-8 encoding to properly handle Bengali text
        df = pd.read_csv(csv_path, encoding='utf-8')
        
        # Convert DataFrame to text representation
        text_parts = []
        
        # Add header information
        text_parts.append("CSV Data Analysis")
        text_parts.append("=" * 50)
        text_parts.append("")
        
        # Add column information
        text_parts.append("📊 Dataset Overview:")
        text_parts.append(f"• Total columns: {len(df.columns)}")
        text_parts.append(f"• Total rows: {len(df)}")
        text_parts.append(f"• File: {os.path.basename(csv_path)}")
        text_parts.append("")
        
        # Add column details
        text_parts.append("📋 Column Information:")
        for i, col in enumerate(df.columns, 1):
            col_type = str(df[col].dtype)
            non_null_count = df[col].count()
            text_parts.append(f"{i}. {col} ({col_type}) - {non_null_count} non-null values")
        text_parts.append("")
        
        # Add sample data with better formatting
        sample_rows = min(5, len(df))
        text_parts.append(f"📄 Sample Data (first {sample_rows} rows):")
        text_parts.append("-" * 40)
        
        for i, row in df.head(sample_rows).iterrows():
            text_parts.append(f"Row {i+1}:")
            for col, val in row.items():
                if pd.isna(val):
                    text_parts.append(f"  {col}: [Empty]")
                else:
                    # Truncate very long values
                    val_str = str(val)
                    if len(val_str) > 100:
                        val_str = val_str[:97] + "..."
                    text_parts.append(f"  {col}: {val_str}")
            text_parts.append("")
        
        # Add summary statistics for numeric columns
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        if len(numeric_cols) > 0:
            text_parts.append("📈 Numeric Column Statistics:")
            text_parts.append("-" * 40)
            for col in numeric_cols:
                stats = df[col].describe()
                text_parts.append(f"{col}:")
                text_parts.append(f"  • Mean: {stats['mean']:.2f}")
                text_parts.append(f"  • Std: {stats['std']:.2f}")
                text_parts.append(f"  • Min: {stats['min']:.2f}")
                text_parts.append(f"  • Max: {stats['max']:.2f}")
                text_parts.append("")
        
        # Add categorical column information
        categorical_cols = df.select_dtypes(include=['object']).columns
        if len(categorical_cols) > 0:
            text_parts.append("🏷️ Categorical Column Information:")
            text_parts.append("-" * 40)
            for col in categorical_cols:
                unique_count = df[col].nunique()
                text_parts.append(f"{col}: {unique_count} unique values")
            text_parts.append("")
        
        return "\n".join(text_parts)
    
    except Exception as e:
        raise Exception(f"Error extracting text from CSV: {str(e)}")

def extract_text_from_db(db_path: str) -> str:
    """Extract text from SQLite database file using pandas and sqlite3 with enhanced formatting"""
    if not PANDAS_AVAILABLE:
        raise Exception("Pandas is not available. Please install pandas to process database files.")
    
    try:
        import sqlite3
        
        # Connect to the database
        conn = sqlite3.connect(db_path)
        
        # Get list of tables
        cursor = conn.cursor()
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
        tables = cursor.fetchall()
        
        if not tables:
            return "Database contains no tables."
        
        text_parts = []
        text_parts.append("🗄️ SQLite Database Analysis")
        text_parts.append("=" * 50)
        text_parts.append(f"📁 Database: {os.path.basename(db_path)}")
        text_parts.append(f"📊 Total tables: {len(tables)}")
        text_parts.append("")
        
        # Analyze each table
        for table in tables:
            table_name = table[0]
            text_parts.append(f"📋 Table: {table_name}")
            text_parts.append("-" * 40)
            
            try:
                # Get table schema
                cursor.execute(f"PRAGMA table_info({table_name})")
                columns = cursor.fetchall()
                
                text_parts.append(f"🔧 Schema ({len(columns)} columns):")
                for i, col in enumerate(columns, 1):
                    col_name, col_type = col[1], col[2]
                    text_parts.append(f"  {i}. {col_name} ({col_type})")
                
                # Get row count
                cursor.execute(f"SELECT COUNT(*) FROM {table_name}")
                row_count = cursor.fetchone()[0]
                text_parts.append(f"📊 Total rows: {row_count}")
                text_parts.append("")
                
                if row_count > 0:
                    # Get sample data (first 3 rows for cleaner output)
                    sample_rows = min(3, row_count)
                    df_sample = pd.read_sql_query(f"SELECT * FROM {table_name} LIMIT {sample_rows}", conn)
                    
                    text_parts.append(f"📄 Sample Data (first {sample_rows} rows):")
                    for i, row in df_sample.iterrows():
                        text_parts.append(f"Row {i+1}:")
                        for col, val in row.items():
                            # Truncate very long values
                            val_str = str(val)
                            if len(val_str) > 80:
                                val_str = val_str[:77] + "..."
                            text_parts.append(f"  {col}: {val_str}")
                        text_parts.append("")
                    
                    # Get summary statistics for numeric columns
                    numeric_cols = df_sample.select_dtypes(include=[np.number]).columns
                    if len(numeric_cols) > 0:
                        text_parts.append("📈 Numeric Column Statistics:")
                        for col in numeric_cols:
                            try:
                                stats = df_sample[col].describe()
                                text_parts.append(f"  {col}:")
                                text_parts.append(f"    • Mean: {stats['mean']:.2f}")
                                text_parts.append(f"    • Std: {stats['std']:.2f}")
                                text_parts.append(f"    • Min: {stats['min']:.2f}")
                                text_parts.append(f"    • Max: {stats['max']:.2f}")
                            except:
                                text_parts.append(f"  {col}: [Statistics not available]")
                        text_parts.append("")
                
            except Exception as table_error:
                text_parts.append(f"❌ Error analyzing table {table_name}: {str(table_error)}")
                text_parts.append("")
        
        conn.close()
        return "\n".join(text_parts)
    
    except Exception as e:
        raise Exception(f"Error extracting text from SQLite database: {str(e)}")

def extract_text_from_file(file_path: str, use_ocr: bool = False) -> str:
    """Extract text from file based on its extension"""
    file_extension = file_path.lower().split('.')[-1]
    
    if file_extension == 'pdf':
        return extract_text_from_pdf(file_path, use_ocr)
    elif file_extension == 'docx':
        return extract_text_from_docx(file_path)
    elif file_extension == 'txt':
        return extract_text_from_txt(file_path)
    elif file_extension == 'csv':
        return extract_text_from_csv(file_path)
    elif file_extension == 'db':
        return extract_text_from_db(file_path)
    else:
        raise Exception(f"Unsupported file type: {file_extension}")

def clean_text(text: str) -> str:
    """Clean and normalize text for better chunking"""
    import re
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove excessive newlines
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    # Clean up bullet points and lists
    text = re.sub(r'^\s*[-•*]\s*', '• ', text, flags=re.MULTILINE)
    
    # Normalize quotes
    text = text.replace('"', '"').replace('"', '"')
    text = text.replace(''', "'").replace(''', "'")
    
    # Remove control characters
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
    
    return text.strip()

def find_semantic_boundary(text: str, position: int, direction: int = -1, max_look: int = 300) -> int:
    """Find the best semantic boundary near the given position"""
    import re
    
    # Priority order for boundaries (from strongest to weakest)
    boundary_patterns = [
        r'\.\s+[A-Z]',      # Sentence endings followed by capital letter
        r'\.\s*\n\s*[A-Z]', # Sentence endings with newline
        r'\.\s+',           # Any sentence ending
        r'!\s+',            # Exclamation marks
        r'\?\s+',           # Question marks
        r';\s+',            # Semicolons
        r',\s+',            # Commas
        r'\n\s*\n',         # Double newlines (paragraph breaks)
        r'\n\s*',           # Single newlines
        r'\s+',             # Any whitespace
    ]
    
    start_pos = max(0, position - max_look) if direction == -1 else position
    end_pos = min(len(text), position + max_look) if direction == 1 else position
    
    search_text = text[start_pos:end_pos]
    
    for pattern in boundary_patterns:
        matches = list(re.finditer(pattern, search_text))
        if matches:
            if direction == -1:
                # Look backwards from position
                for match in reversed(matches):
                    boundary_pos = start_pos + match.end()
                    if boundary_pos <= position:
                        return boundary_pos
            else:
                # Look forwards from position
                for match in matches:
                    boundary_pos = start_pos + match.start()
                    if boundary_pos >= position:
                        return boundary_pos
    
    # If no good boundary found, return the original position
    return position

def chunk_text(text: str, chunk_size: int = 800, overlap: int = 150, min_chunk_size: int = 200) -> List[str]:
    """Split text into clean, meaningful chunks with semantic boundaries and overlap"""
    
    # Clean the text first
    text = clean_text(text)
    
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        # Calculate the ideal end position
        ideal_end = start + chunk_size
        
        if ideal_end >= len(text):
            # Last chunk - take everything remaining
            chunk = text[start:].strip()
            if chunk and len(chunk) >= min_chunk_size:
                chunks.append(chunk)
            break
        
        # Find the best semantic boundary near the ideal end
        semantic_end = find_semantic_boundary(text, ideal_end, direction=-1)
        
        # Ensure we don't go too far back
        min_end = start + (chunk_size // 2)  # At least half the chunk size
        if semantic_end < min_end:
            semantic_end = find_semantic_boundary(text, ideal_end, direction=1)
            if semantic_end > ideal_end + (chunk_size // 2):
                semantic_end = ideal_end
        
        # Extract the chunk
        chunk = text[start:semantic_end].strip()
        
        # Only add chunks that meet minimum size requirements
        if chunk and len(chunk) >= min_chunk_size:
            chunks.append(chunk)
        
        # Calculate next start position with overlap
        next_start = semantic_end - overlap
        
        # Ensure we're making progress
        if next_start <= start:
            next_start = start + (chunk_size // 2)
        
        start = next_start
        
        # Safety check to prevent infinite loops
        if start >= len(text):
            break
    
    # Post-process chunks to ensure quality
    processed_chunks = []
    for chunk in chunks:
        # Remove chunks that are too short or contain mostly whitespace
        if len(chunk.strip()) >= min_chunk_size:
            # Clean up the chunk
            cleaned_chunk = clean_text(chunk)
            if cleaned_chunk:
                processed_chunks.append(cleaned_chunk)
    
    return processed_chunks

class DocumentRAG:
    def __init__(self, storage_file='documents.json'):
        self.storage_file = storage_file
        # Use TF-IDF instead of sentence transformers for simplicity
        self.vectorizer = TfidfVectorizer(
            stop_words='english',
            max_features=1000,
            ngram_range=(1, 2)
        )
        self.documents = []
        self.embeddings = None
        self.is_fitted = False
        
        # Load existing documents from file
        self._load_documents()
        
    def _load_documents(self):
        """Load documents from storage file"""
        if os.path.exists(self.storage_file):
            try:
                with open(self.storage_file, 'r', encoding='utf-8') as f:
                    self.documents = json.load(f)
                print(f"Loaded {len(self.documents)} documents from {self.storage_file}")
                if self.documents:
                    self._update_embeddings()
            except Exception as e:
                print(f"Error loading documents: {e}")
                self.documents = []
        else:
            print(f"No existing documents found. Starting fresh.")
            self.documents = []
    
    def _save_documents(self):
        """Save documents to storage file"""
        try:
            with open(self.storage_file, 'w', encoding='utf-8') as f:
                json.dump(self.documents, f, indent=2, ensure_ascii=False)
            print(f"Saved {len(self.documents)} documents to {self.storage_file}")
        except Exception as e:
            print(f"Error saving documents: {e}")
        
    def add_document(self, text: str, metadata: Dict[str, Any] = None):
        """Add a document to the knowledge base"""
        if metadata is None:
            metadata = {}
        
        # Store document with metadata
        doc = {
            'text': text,
            'metadata': metadata,
            'id': len(self.documents)
        }
        self.documents.append(doc)
        
        # Re-fit the vectorizer with all documents
        self._update_embeddings()
        
        # Save to file
        self._save_documents()
        
        return doc['id']
    
    def add_file_document(self, file_path: str, metadata: Dict[str, Any] = None, use_ocr: bool = False) -> List[int]:
        """Add a document file (PDF, DOCX, or TXT) to the knowledge base by extracting and chunking text"""
        if metadata is None:
            metadata = {}
        
        # Extract text from file
        full_text = extract_text_from_file(file_path, use_ocr)
        
        # Chunk the text
        chunks = chunk_text(full_text)
        
        # Add each chunk as a separate document
        doc_ids = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata.update({
                'source': 'file_upload',
                'file_type': file_path.lower().split('.')[-1],
                'chunk_id': i,
                'total_chunks': len(chunks),
                'filename': os.path.basename(file_path),
                'ocr_used': use_ocr
            })
            
            doc_id = self.add_document(chunk, chunk_metadata)
            doc_ids.append(doc_id)
        
        return doc_ids
    
    def _update_embeddings(self):
        """Update embeddings when documents change"""
        if not self.documents:
            return
        
        # Extract all document texts
        texts = [doc['text'] for doc in self.documents]
        
        # Fit vectorizer and create embeddings
        self.embeddings = self.vectorizer.fit_transform(texts)
        self.is_fitted = True
    
    def search(self, query: str, top_k: int = 3):
        """Search for relevant documents using TF-IDF similarity"""
        if not self.documents or not self.is_fitted:
            return []
        
        # Transform query to vector
        query_vector = self.vectorizer.transform([query])
        
        # Calculate similarities
        similarities = cosine_similarity(query_vector, self.embeddings)[0]
        
        # Get top-k most similar documents
        top_indices = np.argsort(similarities)[::-1][:top_k]
        
        results = []
        for idx in top_indices:
            if similarities[idx] > 0:  # Only include documents with some similarity
                results.append({
                    'document': self.documents[idx],
                    'similarity': float(similarities[idx])
                })
        
        return results
    
    def generate_response(self, query: str, top_k: int = 3):
        """Generate a response using retrieved context"""
        # Search for relevant documents
        search_results = self.search(query, top_k)
        
        if not search_results:
            return {
                'response': 'I don\'t have enough information to answer that question.',
                'context': [],
                'query': query
            }
        
        # Build context from retrieved documents
        context = []
        for result in search_results:
            context.append({
                'text': result['document']['text'],
                'similarity': result['similarity'],
                'metadata': result['document']['metadata']
            })
        
        # Simple response generation (in a real system, you'd use an LLM here)
        # For this demo, we'll just return the most relevant document
        best_match = context[0]
        
        response = f"Based on the most relevant information I found:\n\n{best_match['text']}\n\n(Similarity score: {best_match['similarity']:.3f})"
        
        return {
            'response': response,
            'context': context,
            'query': query
        }

# Initialize RAG system with document support
rag = DocumentRAG()

# Add sample documents only if no documents exist
if len(rag.documents) == 0:
    print("Adding sample documents...")
    sample_docs = [
        {
            'text': 'RAG stands for Retrieval-Augmented Generation. It combines information retrieval with text generation to provide more accurate and contextual responses.',
            'metadata': {'topic': 'definition', 'source': 'educational'}
        },
        {
            'text': 'Vector embeddings are numerical representations of text that capture semantic meaning. They allow us to find similar documents by comparing their embeddings.',
            'metadata': {'topic': 'embeddings', 'source': 'technical'}
        },
        {
            'text': 'The RAG pipeline typically involves: 1) Storing documents with embeddings, 2) Searching for relevant documents when given a query, 3) Using retrieved context to generate a response.',
            'metadata': {'topic': 'pipeline', 'source': 'process'}
        },
        {
            'text': 'Cosine similarity is a common method for comparing vector embeddings. It measures the cosine of the angle between two vectors, ranging from -1 to 1.',
            'metadata': {'topic': 'similarity', 'source': 'mathematical'}
        },
        {
            'text': 'TF-IDF stands for Term Frequency-Inverse Document Frequency. It is a statistical measure used to evaluate how important a word is to a document in a collection.',
            'metadata': {'topic': 'tfidf', 'source': 'statistical'}
        }
    ]
    
    # Add sample documents to the RAG system
    for doc in sample_docs:
        rag.add_document(doc['text'], doc['metadata'])

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'message': 'Document-Enabled RAG API is running',
        'documents_count': len(rag.documents),
        'method': 'TF-IDF with PDF, DOCX, TXT, CSV, and SQLite DB support',
        'storage_file': rag.storage_file,
        'supported_formats': list(ALLOWED_EXTENSIONS),
        'ocr_available': OCR_AVAILABLE,
        'ocr_note': 'OCR requires Tesseract to be installed on the system'
    })

@app.route('/documents', methods=['GET'])
def get_documents():
    """Get all documents in the knowledge base"""
    return jsonify({
        'documents': rag.documents,
        'count': len(rag.documents)
    })

@app.route('/documents', methods=['POST'])
def add_document():
    """Add a new document to the knowledge base"""
    data = request.get_json()
    
    if not data or 'text' not in data:
        return jsonify({'error': 'Text field is required'}), 400
    
    text = data['text']
    metadata = data.get('metadata', {})
    
    doc_id = rag.add_document(text, metadata)
    
    return jsonify({
        'message': 'Document added successfully and saved to file',
        'document_id': doc_id,
        'total_documents': len(rag.documents),
        'storage_file': rag.storage_file
    }), 201

@app.route('/upload-document', methods=['POST'])
def upload_document():
    """Upload and process a document file (PDF or DOCX)"""
    # Check if file was uploaded
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    # Check if file was selected
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    # Check if file type is allowed
    if not allowed_file(file.filename):
        return jsonify({'error': f'File type not allowed. Supported formats: {", ".join(ALLOWED_EXTENSIONS)}'}), 400
    
    try:
        # Save the uploaded file
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Get metadata from form data
        metadata = {}
        if request.form.get('topic'):
            metadata['topic'] = request.form.get('topic')
        if request.form.get('source'):
            metadata['source'] = request.form.get('source')
        
        # Check if OCR should be used
        use_ocr = request.form.get('use_ocr', 'false').lower() == 'true'
        
        # Process the document
        doc_ids = rag.add_file_document(filepath, metadata, use_ocr)
        
        # Clean up the uploaded file
        os.remove(filepath)
        
        return jsonify({
            'message': 'Document processed successfully',
            'filename': filename,
            'file_type': filename.lower().split('.')[-1],
            'document_ids': doc_ids,
            'chunks_created': len(doc_ids),
            'total_documents': len(rag.documents),
            'storage_file': rag.storage_file
        }), 201
        
    except Exception as e:
        # Clean up file if it exists
        if os.path.exists(filepath):
            os.remove(filepath)
        return jsonify({'error': f'Error processing document: {str(e)}'}), 500

@app.route('/search', methods=['POST'])
def search_documents():
    """Search for relevant documents"""
    data = request.get_json()
    
    if not data or 'query' not in data:
        return jsonify({'error': 'Query field is required'}), 400
    
    query = data['query']
    top_k = data.get('top_k', 3)
    
    results = rag.search(query, top_k)
    
    return jsonify({
        'query': query,
        'results': results,
        'count': len(results)
    })

@app.route('/generate', methods=['POST'])
def generate_response():
    """Generate a response using RAG"""
    data = request.get_json()
    
    if not data or 'query' not in data:
        return jsonify({'error': 'Query field is required'}), 400
    
    query = data['query']
    top_k = data.get('top_k', 3)
    
    response = rag.generate_response(query, top_k)
    
    return jsonify(response)

@app.route('/')
def home():
    """Home page with API documentation"""
    return jsonify({
        'message': 'Welcome to the Document-Enabled RAG API!',
        'note': 'This version supports PDF, DOCX, TXT, CSV, and SQLite DB uploads with optional OCR for image-based PDFs',
        'storage_file': rag.storage_file,
        'endpoints': {
            'GET /health': 'Health check',
            'GET /documents': 'Get all documents',
            'POST /documents': 'Add a new document',
            'POST /upload-document': 'Upload and process a document file (PDF/DOCX/TXT/CSV/DB)',
            'POST /search': 'Search for relevant documents',
            'POST /generate': 'Generate a response using RAG'
        },
        'example_usage': {
            'add_document': {
                'method': 'POST',
                'endpoint': '/documents',
                'body': {'text': 'Your document text', 'metadata': {'topic': 'example'}}
            },
            'upload_document': {
                'method': 'POST',
                'endpoint': '/upload-document',
                'form_data': {'file': 'your_file.pdf', 'topic': 'example', 'source': 'upload'}
            },
            'search': {
                'method': 'POST',
                'endpoint': '/search',
                'body': {'query': 'What is RAG?', 'top_k': 3}
            },
            'generate': {
                'method': 'POST',
                'endpoint': '/generate',
                'body': {'query': 'How does RAG work?', 'top_k': 3}
            }
        }
    })

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5004) 